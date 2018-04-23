# file handling
import codecs
import os
from os.path import dirname
import shutil

# webscraping assets
import requests
import json

# sorting
import re

# janome
from janome.tokenizer import Tokenizer
from janome.analyzer import Analyzer
from janome.charfilter import *
from janome.tokenfilter import *

# export to docx
from docx import Document
from docx.shared import Mm

# opens a specific file and returns its content in a string
def open(filename):
	file = codecs.open(filename,"r","utf-8")
	raw_string = file.read()
	file.close()	
	return raw_string

# this function helps to sort files with complex filenames
def numericalSort(filelist):
	numbers = re.compile(r'(\d+)')
	parts = numbers.split(filelist)
	parts[1::2] = map(int, parts[1::2])
	return parts

# sorts and returns locations for all vocab and kanji files of the desired grading system
# JLPT files are sorted in inverse order, all other files in numerical order
def reflists(filename):
	print("\tCompiling reference data for "+filename.split("_")[1]+' grading system...')
	vocab_files = []
	kanji_files = []
	ref_files = sorted(os.listdir(dirname(os.path.abspath(""))+"\data"),key=numericalSort)
	for file in ref_files:
		if filename.split(".")[0].split("_")[1] in file:			
			if "kanji" in file:							
				if file.split("_")[0] == "JLPT":			
					kanji_files.insert(0,file)   
				else:
					kanji_files.append(file)			
			if "vocab" in file:					
				if file.split("_")[0] == "JLPT":	
					vocab_files.insert(0,file)
				else:
					vocab_files.append(file)			
	print("\t\tVocab files:",vocab_files[0],"->",vocab_files[-1])
	print("\t\tKanji files:",kanji_files[0],"->",kanji_files[-1])	
	print("\tDone.")
	return vocab_files, kanji_files

# Janome tokenizes the text
# punctuation and other unused characters are removed in the regex filter
# particles, verb endings and other unused parts of speech are suppressed by the POSStopFilter
def janome(raw_string):
	print("\tTokenising your file...")
	char_filters = [UnicodeNormalizeCharFilter(), RegexReplaceCharFilter(u'[・.?!:; <>+*/\-()%\'\"~•❶❷°&]|[0-9]|[a-zA-Z]', u'\b')]
	tokenizer = Tokenizer()
	token_filters = [POSStopFilter(["助詞","助動詞","記号"]), LowerCaseFilter()] 
	a = Analyzer(char_filters, tokenizer, token_filters)
	token_dict = {}
	alternatives = ""
	level = ""
	# Janome export format: ["表層形,	品詞, 品詞細分類1, 品詞細分類2, 品詞細分類3, 活用型, 活用形, 基本形, 読み, 表層形", ""]
	#						["surface	PoS, PoS_detail_1, PoS_detail_2, PoS_detail_3, infl_type, infl_form, lemma, reading, phonetic", ""]
	for token_info in a.analyze(raw_string):	
		lemma = str(token_info).split(",")[6]
		reading = str(token_info).split(",")[7] 
		if str(token_info).split(",")[1] == "固有名詞":
			PoS = "proper name"
			meaning = "• proper name"
		else:
			PoS = ""
			meaning = ""
		token_dict[lemma] = [alternatives, PoS, level, reading, meaning]
		# print("\t",lemma, token_dict[lemma])
	print("\tDone.")									
	return token_dict

# looks up and injects the level of every item into the dictionary
# the level returned is the level on which the wordkanji is learned
def check_level(token_dict, vocab_files, kanji_files):
	print("\tChecking the level of every token...")
	counter = 0
	hiragana = "ぁあぃいぅうぇえぉおかがきぎくぐけげこごさざしじすずせぜそぞただちぢっつづてでとどなにぬねのはばぱひびぴふぶぷへべぺほぼぽまみむめもゃやゅゆょよらりるれろゎわゐをんゕゖ"
	katakana = "ァアィイゥウェエォオカガキギクグケゲコゴサザシジスズセゼソゾタダチヂッツヅテデトドナニヌネノハバパヒビピフブプヘベペホボポマミムメモャヤュユョヨラリルレロヮワヲンヴヵヶヷヺー"
	kana = hiragana + katakana
	for token in token_dict:
		counter = 0
		level = ""
		while level == "":
			try:
				if token in (open(dirname(os.path.abspath(""))+"\data/"+vocab_files[counter])): 				# checks the first vocab file for the item
					level = counter	
					token_dict[token][2] = vocab_files[level].split("_")[1]# using the file list to inject the right level name
					print("\t\t",token,"found in vocab files: assigning level",vocab_files[level].split("_")[1])
				else:															# if not found, the next file is checked
					counter += 1 
			except IndexError:											# returns "unknown" if not found in any vocablist
				counter = 0
				while level == "":
					try:
						if token in (open(dirname(os.path.abspath(""))+"\data/"+kanji_files[counter])): 				# checks the first vocab file for the item
							level = counter	
							token_dict[token][2] = kanji_files[level].split("_")[1]# using the file list to inject the right level name
							print("\t\t",token,"found in kanji files: assigning level",kanji_files[level].split("_")[1])
						else:															# if not found, the next file is checked
							counter += 1 
					except IndexError:	
						level = "unknown"			# returns "unknown" if not found in any vocablist
						token_dict[token][2] = level
		if token_dict[token][2] == "unknown":							# for all failed searches, splits the lemma into components
			print("\t\t",token,"not found. Looking up individual components.")
			highest = 0		
			for character in token:
				if character in kana: 
					character_level = 0
					if character_level >= highest:
						token_dict[token][2] = vocab_files[highest].split("_")[1]
					print("\t\t\t",character,"found in kana. Added the minimum level:",vocab_files[character_level].split("_")[1])
				else:
					character_level = ""
					counter = 0
					while character_level == "":
						try:
							if character in open(dirname(os.path.abspath(""))+"\data/"+kanji_files[counter]):
								character_level = counter
								print("\t\t\t",character,"found in kanji file",kanji_files[counter].split("_")[1])
								if character_level > highest:
									highest = counter
								token_dict[token][2] = kanji_files[highest].split("_")[1]
							else: 
								counter += 1
						except IndexError:
							print("\t\t\t",character,"not found in kanji files. Checking vocab files")
							counter = 0
							while character_level == "":	
								try:
									if character in open(dirname(os.path.abspath(""))+"\data/"+vocab_files[counter]):
										character_level = counter
										print("\t\t\t",character,"found in vocab file",vocab_files[counter].split("_")[1])
										if character_level > highest:
											highest = counter
											token_dict[token][2] = vocab_files[highest].split("_")[1]
									else: 
										counter += 1
								except IndexError:
									character_level = counter
									if character_level > highest:
										highest = counter - 1
									print("\t\t\t",character,"is above level",vocab_files[highest].split("_")[1])
									token_dict[token][2] = vocab_files[highest].split("_")[1]
			print("\t\t\tAssigned the highest found level:", token_dict[token][2])
		# print(token, token_dict[token])
	print("\tDone.")
	return token_dict

# this function determines the amount of words per level for all dictionary items
def text_level(token_dict, filename):
	print("\tCalculating and appending the recommended",filename.split(".")[0].split("_")[1],"level.")
	level_dict = {}
	for item in token_dict:
		if token_dict[item][2] not in level_dict:
			level_dict[token_dict[item][2]] = 1
		else: 
			level_dict[token_dict[item][2]] += 1
	sorted_keys = sorted(level_dict, key=level_dict.get, reverse = True)
	new_name = filename.split(".")[0]+"_"+sorted_keys[0]+".txt"
	shutil.move(dirname(os.path.abspath(""))+"\input/"+filename,dirname(os.path.abspath(""))+"\input/"+new_name)
	print("\t\tLevel:",sorted_keys[0],">",sorted_keys[1],">",sorted_keys[2])
	print("\t\tItems:",sorted(level_dict.values(), reverse = True)[0],">",sorted(level_dict.values(), reverse = True)[1],">",sorted(level_dict.values(), reverse = True)[2])
	print("\t\tNew filename:",new_name)
	print("\tDone.")
	return

# keeps only the items that have a higher (harder) level than the input level
def filter_dict(token_dict, filename):
	print("\tDetected level filter. Filtering items above level",filename.split(".")[0].split("_")[1],filename.split(".")[0].split("_")[2]+"...")
	gloss = {}
	if filename.split("_")[1] == "JLPT":
		for item in token_dict:
			if int(token_dict[item][2][1]) <= int(filename.split(".")[0].split("_")[2][1]):
				gloss[item] = token_dict[item]
				# print("\t",item,token_dict[item])
	else: 
		for item in token_dict:
			if int(token_dict[item][2]) >= int(filename.split(".")[0].split("_")[2]):
				gloss[item] = token_dict[item]
				# print("\t",item,token_dict[item])
	print("\tDone.")
	return gloss

# sends queries to the Jisho.org API and injects useful information from the JSON output into the dictionary
def jisho_api_lookup(gloss):
	print("\tInjecting data from Jisho.org...")
	for query in gloss:
		r = requests.get('http://jisho.org/api/v1/search/words?keyword="'+query+'"') 
		data = json.loads(r.text)
		words = []
		readings = []
		alternatives = [query]
		definitions = []	
		PoS = []
		gloss[query][0] = query
		# Start of JSON navigation
		for resultdict in range(len(data["data"])):
			for JP_result in range(len(data["data"][resultdict]["japanese"])):
				try:
					# if alternative ways of writings are found, appends them between brackets to the lemma.
					if data["data"][resultdict]["japanese"][JP_result]["word"] != query and "(" not in data["data"][resultdict]["japanese"][JP_result]["word"]:
						alternatives.append(data["data"][resultdict]["japanese"][JP_result]["word"])
						gloss[query][0] = query+" ("+", ".join(alternatives[1:3])+")"
				except:
					# words unknown to the API will return no "word" attribute or value
					# values from the janome function are kept; usually proper names
					pass
				try:
					# finds the possible pronunciations of the lemma and appends them
					if data["data"][resultdict]["japanese"][JP_result]["reading"] not in readings:
						readings.append(data["data"][resultdict]["japanese"][JP_result]["reading"])
						gloss[query][3] = "\n".join(readings[0:2]) # this index modifies the amount of meanings
				except:
					# words unknown to the API will return no "reading" attribute or value
					# values from the janome function are kept; usually proper names
					pass
			for sense_result in range(len(data["data"][resultdict]["senses"])):
				for meaning in range(len(data["data"][resultdict]["senses"][sense_result]["english_definitions"])):
					# definitions are checked for duplicates, and injected into the dictionary
					if "• "+", ".join(data["data"][resultdict]["senses"][sense_result]["english_definitions"]) not in definitions:
						definitions.append("• "+", ".join(data["data"][resultdict]["senses"][sense_result]["english_definitions"]))
						gloss[query][4] = "\n".join(definitions[0:5]) # this index modifies the amount of definitions
					# possible parts of speech are checked for duplicates, and injected into the dictionary
				for part_of_speech in range(len(data["data"][resultdict]["senses"][sense_result]["parts_of_speech"])):
					if data["data"][resultdict]["senses"][sense_result]["parts_of_speech"][0] not in PoS and "Wikipedia definition" not in data["data"][resultdict]["senses"][sense_result]["parts_of_speech"]:
						PoS.append(data["data"][resultdict]["senses"][sense_result]["parts_of_speech"][0])
						gloss[query][1] = "\n".join(PoS) 
		print("\t\t",query,gloss[query])
	print("\tDone.")
	return gloss

# docx table column width has to be defined for every single cell
def set_column_width(column, width):
	for cell in column.cells:
		cell.width = width
	return

# the glossary is exported to a pretty docx file
def word_export(gloss, file, raw_text):
	print("\tCreating a docx file...")
	document = Document()
	document.add_heading(file.split(".")[0].split("_")[0]+" ("+file.split(".")[0].split("_")[1]+" "+file.split(".")[0].split("_")[2]+")", 0)
	paragraph = document.add_paragraph()
	table = document.add_table(rows=1,cols=3,style='Table Grid')
	table.autofit = False
	table.rows[0].cells[0].text = "Word"
	table.rows[0].cells[1].text = "Reading"
	table.rows[0].cells[2].text = "Meaning"
	for line in gloss:
		cells = table.add_row().cells
		cells[0].text = str(gloss[line][0])
		cells[1].text = str(gloss[line][3])
		cells[2].text = str(gloss[line][4][0:])
	set_column_width(table.columns[0], Mm(45))
	set_column_width(table.columns[1], Mm(35))
	set_column_width(table.columns[2], Mm(100))
	sections = document.sections
	for section in sections:
		section.top_margin = Mm(19)
		section.bottom_margin = Mm(19)
		section.left_margin = Mm(19)
		section.right_margin = Mm(19)
	document.save(dirname(os.path.abspath(""))+"/"+file.split(".")[0].split("_")[0]+"_glossary_"+file.split(".")[0].split("_")[1]+"_"+file.split(".")[0].split("_")[2]+".docx")	
	print("\t\tCreated file",file.split(".")[0].split("_")[0]+"_glossary_"+file.split(".")[0].split("_")[1]+"_"+file.split(".")[0].split("_")[2]+".docx")
	print("\tDone.")
	return

# executes the main script
def glossary_extractor():
	for file in os.listdir(dirname(os.path.abspath(""))+"\input"):
		print("Processing "+file+"...")
		vocab_files, kanji_files = reflists(file)
		raw_text = open(dirname(os.path.abspath(""))+"\input\/"+file)
		janome_dict = janome(raw_text)
		leveled_dict = check_level(janome_dict, vocab_files, kanji_files)
		if len(file.split("_"))<3:
			text_level(leveled_dict, file)
			print("Your file has been modified successfully.")
		else:
			gloss = filter_dict(leveled_dict, file)
			completed_gloss = jisho_api_lookup(gloss)
			word_export(completed_gloss, file, raw_text)
			print("Your glossary has been generated successfully.")
	return 


glossary_extractor()


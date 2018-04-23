import codecs
import json
import os
import numpy
import csv
import xlsxwriter
from collections import OrderedDict
from operator import itemgetter
from docx import Document


def overviewlex(project, type_pr):
	if type_pr == 'Python_corp':
		pad_stat = os.path.join('..', 'output', 'stats', project)
		pad_view = os.path.join('..', 'output', 'view', project)
	if type_pr == 'Local_project':
		pad_stat = os.path.join(project, 'metadata')
		pad_view = os.path.join(project, 'view')
	if os.path.isdir(pad_view) == False:
		os.mkdir(pad_view)
	g = codecs.open(os.path.join(pad_stat, project+'_freq_lex.json'), 'r', 'utf-8-sig')
	d_lex = json.load(g)
	g.close()
	##
	##
	# d_lex = json.load(g)
	# g.close()
	d_items = {}
	for cat in ['V', 'NC', 'ADJ', 'ADJV']:
		d_items_cat = {}
		adv_spec_high = []
		adv_spec_medium = []
		adv_spec_low = []
		adv_spec_none = []
		all_spec_high = []
		all_spec_medium = []
		all_spec_low = []
		for lem in d_lex[cat]:
			dif = []
			for crit in d_lex[cat][lem]:
				if crit.startswith('dif_'):
					dif.append(d_lex[cat][lem][crit])
			if d_lex[cat][lem]['lev'] == 0 and 'A' in dif: #overwegen om dit toch te kwantificeren in cijfer? vervangen door formule van diff?
				adv_spec_high.append(lem)
			if d_lex[cat][lem]['lev'] == 0 and any(level in dif for level in ['A', 'B']): #overwegen om dit toch te kwantificeren in cijfer? vervangen door formule van diff?
				adv_spec_medium.append(lem)
			if d_lex[cat][lem]['lev'] == 0 and 'D' not in dif: #overwegen om dit toch te kwantificeren in cijfer? vervangen door formule van diff?
				adv_spec_low.append(lem)
			if d_lex[cat][lem]['lev'] == 0: #overwegen om dit toch te kwantificeren in cijfer? vervangen door formule van diff?
				adv_spec_none.append(lem)
			if 'A' in dif: #overwegen om dit toch te kwantificeren in cijfer? vervangen door formule van diff?
				all_spec_high.append(lem)
			if any(level in dif for level in ['A', 'B']):
				all_spec_medium.append(lem)
			if 'D' not in dif: #overwegen om dit toch te kwantificeren in cijfer? vervangen door formule van diff?
				all_spec_low.append(lem)
		d_items_cat['adv_spec_high'] = adv_spec_high
		d_items_cat['adv_spec_medium'] = adv_spec_medium
		d_items_cat['adv_spec_low'] = adv_spec_low
		d_items_cat['adv_spec_none'] = adv_spec_none
		d_items_cat['all_spec_high'] = all_spec_high
		d_items_cat['all_spec_medium'] = all_spec_medium
		d_items_cat['all_spec_low'] = all_spec_low
		d_items[cat] = d_items_cat
	g = codecs.open(os.path.join(pad_stat, project + '_class_lex.json'), 'w', 'utf-8-sig') #hebben we dit bestand nodig?
	json.dump(d_items, g)
	g.close()
	wb = xlsxwriter.Workbook(os.path.join(pad_view, project+'_voc_overview.xlsx'))
	bold = wb.add_format({'bold': True})
	ws = wb.add_worksheet('overview')
	ws.autofilter('A1:B1')
	col = 2
	ws.write(0, 0, 'Level', bold)
	ws.write(0, 1, 'Specificity', bold)
	ws.write(0, 2, 'V', bold)
	ws.write(0, 3, 'NC', bold)	
	ws.write(0, 4, 'ADJ', bold)
	ws.write(0, 5, 'ADJV', bold)
	ws.write(0, 6, 'TOTAL', bold)
	ws.write(1, 0, 'advanced', bold)
	ws.write(2, 0, 'advanced', bold)
	ws.write(3, 0, 'advanced', bold)
	ws.write(4, 0, 'advanced', bold)
	ws.write(5, 0, 'all', bold)
	ws.write(6, 0, 'all', bold)
	ws.write(7, 0, 'all', bold)
	ws.write(1, 1, 'high', bold)
	ws.write(2, 1, 'medium', bold)
	ws.write(3, 1, 'low', bold)
	ws.write(4, 1, 'none', bold)
	ws.write(5, 1, 'high', bold)
	ws.write(6, 1, 'medium', bold)
	ws.write(7, 1, 'low', bold)
	c_adv_spec_high = 0
	c_adv_spec_medium = 0
	c_adv_spec_low = 0
	c_adv_spec_none = 0
	c_all_spec_high = 0
	c_all_spec_medium = 0
	c_all_spec_low = 0
	for cat in ['V', 'NC', 'ADJ', 'ADJV']:
		c_adv_spec_high += len(d_items[cat]['adv_spec_high'])
		c_adv_spec_medium += len(d_items[cat]['adv_spec_medium'])
		c_adv_spec_low += len(d_items[cat]['adv_spec_low'])
		c_adv_spec_none += len(d_items[cat]['adv_spec_none'])
		c_all_spec_high += len(d_items[cat]['all_spec_high'])
		c_all_spec_medium += len(d_items[cat]['all_spec_medium'])
		c_all_spec_low += len(d_items[cat]['all_spec_low'])
		row = 1
		ws.write(row, col, len(d_items[cat]['adv_spec_high']))
		ws.write(row+1, col, len(d_items[cat]['adv_spec_medium']))
		ws.write(row+2, col, len(d_items[cat]['adv_spec_low']))
		ws.write(row+3, col, len(d_items[cat]['adv_spec_none']))
		ws.write(row+4, col, len(d_items[cat]['all_spec_high']))
		ws.write(row+5, col, len(d_items[cat]['all_spec_medium']))
		ws.write(row+6, col, len(d_items[cat]['all_spec_low']))
		ws.write(row, 6, c_adv_spec_high)
		ws.write(row+1, 6, c_adv_spec_medium)
		ws.write(row+2, 6, c_adv_spec_low)
		ws.write(row+3, 6, c_adv_spec_none)
		ws.write(row+4, 6, c_all_spec_high)
		ws.write(row+5, 6, c_all_spec_medium)
		ws.write(row+6, 6, c_all_spec_low)
		col += 1
	ws.set_column(0, 0, len("adv_spec_medium")+4)
	ws.set_column(1, 1, len("adv_spec_medium")+4)
	##
	type_dif = []
	for lem in d_lex['NC']:
		if lem.startswith('a'):
			for dif in d_lex['NC'][lem]:
				if dif.startswith('dif_') and dif.replace('dif_', '') not in type_dif:
					type_dif.append(dif.replace('dif_', ''))
	for cat in ['NC', 'V', 'ADJ', 'ADJV', 'ADV']: #blijven nog over PE en XP
		l = OrderedDict(reversed(sorted(d_lex[cat].items(), key=lambda t: t[1]['freq'])))
		ws = wb.add_worksheet(cat)
		ws.autofilter('A1:E1')
		len_0 = []
		row = 0
		col = 3
		ws.write(row, 0, 'LEMA', bold)
		ws.write(row, 1, 'LEVEL', bold)
		ws.write(row, 2, 'FREQ', bold)
		ws.write(row, 3, 'PCTIL', bold)
		for dif in type_dif:
			ws.write(row, col+1, 'SPEC_'+dif, bold)
		col = 3
		for lem in l:
			row += 1
			ws.write(row, 0, lem)
			len_0.append(len(lem))
			if l[lem]['lev'] == 0:
				ws.write(row, 1, 'avanzado')
			elif l[lem]['lev'] == 1:
				ws.write(row, 1, 'intermedio')
			elif l[lem]['lev'] == 2:
				ws.write(row, 1, 'básico')
			ws.write(row, 2, l[lem]['freq'])
			ws.write(row, 3, l[lem]['pctil'])
			for dif in type_dif:
				ws.write(row, col+1, l[lem]['dif_'+dif])
		ws.set_column(0, 0, max(len_0))
		ws.set_column(1, 1, len("intermedio")+2)
		ws.set_column(3, 3, len("dif_dture")+6)
	wb.close()

def selecttext(project, type_pr, select_level, select_spec, select_pos): 
	print("Selecting most efficient reading order of texts")

	if type_pr == 'Python_corp':
		pad_in_stat = os.path.join('..', 'output', 'stats', project)
		pad_in_sent = os.path.join('..', 'corpus', 'sent', project)
		pad_out_sent_select = os.path.join('..', 'output', 'stats', project)
		pad_out_sent_print = os.path.join('..', 'output', 'view', project)
		pad_index = os.path.join('..', 'corpus', 'metadata', 'index_lem')
	if type_pr == 'Local_project':
		pad_in_stat = os.path.join(project, 'metadata')
		pad_in_sent = os.path.join(project, 'corp_data', 'sent')
		pad_out_sent_select = os.path.join(project, 'metadata')
		pad_out_sent_print = os.path.join(project, 'view')
		pad_index = os.path.join(project, 'metadata')

	g = codecs.open(os.path.join(pad_out_sent_select, project + '_class_lex.json'), 'r', 'utf-8-sig')
	d_lex = json.load(g)
	g.close()
	g = codecs.open(os.path.join(pad_index, 'corp_index_'+project + '.json'), 'r', 'utf-8-sig')
	d_index = json.load(g)
	g.close()
	g = codecs.open(os.path.join(pad_in_stat, project+'_stats.json'), 'r', 'utf-8-sig')
	d_stats = json.load(g)
	g.close()
	if os.path.isdir(pad_out_sent_print) == False:
		os.mkdir(pad_out_sent_print)
	if select_level == 'advanced':
		out_1 = '_lev_adv'
		type_se_1 = 'adv'
		type_se_expl_1 = 'advanced vocabulary, '
	elif select_level == 'all':
		type_se_1 = 'all'
		type_se_expl_1 = 'no restrictions on learning level, '
		out_1 = 'text_all'
	if select_spec == 'high':
		out_2 = '_spec_high'
		type_se_2 = '_high'
		type_se_expl_2 = 'high specificity and high frequency'
	elif select_spec == 'medium':
		type_se_2 = '_spec_medium'
		type_se_expl_2 = 'not very frequent, but more frequent than in the reference point'
		out_2 = '_spec_medium'
	elif select_spec == 'low':
		type_se_2 = '_spec_low'
		type_se_expl_2 = 'non-typical lexemes excluded'
		out_2 = '_spec_low'
	elif select_spec == '':
		type_se_2 = '_spec_none'
		type_se_expl_2 = 'no restrictions on specificity'
		out_2 = '_spec_none'
	type_se = type_se_1+type_se_2
	type_se_expl = type_se_expl_1+type_se_expl_2
	outputfile = out_1 + out_2
	select_pos_text = '_'.join(select_pos)
	list_search= {}
	list_found = {}
	list_texts = []
	list_found_len = {}
	list_texts_len = []
	list_found_comp = {}
	n_words = 0
	n_words_len = 0
	result = ''
	for cat in select_pos:
		list_search[cat] = []
		list_found[cat] = []
		list_found_len[cat] = []
		list_found_comp[cat] = []
		for lem in d_lex[cat][type_se]:
			list_search[cat].append(lem)
	## OPNIEUW ZICHTBAAR MAKEN!
	for c in range(1, 200):
		d_texts = {}
		for cat in select_pos:
			for lem in d_lex[cat][type_se]:
				if lem in list_search[cat] and lem not in list_found[cat]:
					for text in d_index[lem][cat]:
						if text not in list_texts:
							if text not in d_texts:
								d_texts[text] = [(lem, cat)]
							elif (lem, cat) not in d_texts[text]: #
								d_texts[text].append((lem, cat))
		l_pct = []
		d_pct = {}
		l_len = []
		d_len = {}
		for text in d_texts:
			pct = int(len(d_texts[text])/d_stats[text]['c_words']*100)
			l_pct.append(pct)
			d_pct[pct] = text
			l_len.append(len(d_texts[text]))
			d_len[len(d_texts[text])] = text
		if l_pct:
			t_max = d_pct[max(l_pct)]
			list_texts.append(t_max)
			n_words += d_stats[t_max]['c_words']
			for lem in d_texts[t_max]:
				if lem[0] not in list_found[lem[1]]:
					list_found[lem[1]].append(lem[0])
			n_found = 0
			n_search = 0
			for cat in list_found:
				n_found += len(list_found[cat])
			for cat in list_search:
				n_search += len(list_search[cat])
			result += str(c) + '\t' + t_max + '\t' + str(int(n_found/n_search*100)) + '\t' + '%' + '\t' + str(n_words) + '\t' + 'words' + '\n\n\n'
	for c in range(1, 200):
		d_texts = {}
		for cat in select_pos:
			for lem in d_lex[cat][type_se]:
				if lem in list_search[cat] and lem not in list_found_len[cat]:
					for text in d_index[lem][cat]:
						if text not in list_texts_len:
							if text not in d_texts:
								d_texts[text] = [(lem, cat)]
							elif (lem, cat) not in d_texts[text]: #
								d_texts[text].append((lem, cat))
		l_pct = []
		d_pct = {}
		l_len = []
		d_len = {}
		for text in d_texts:
			l_len.append(len(d_texts[text]))
			d_len[len(d_texts[text])] = text
		if l_len:
			t_max = d_len[max(l_len)]
			list_texts_len.append(t_max)
			n_words_len += d_stats[t_max]['c_words']
			for lem in d_texts[t_max]:
				if lem[0] not in list_found_len[lem[1]]:
					list_found_len[lem[1]].append(lem[0])
			n_found = 0
			n_search = 0
			for cat in list_found_len:
				n_found += len(list_found_len[cat])
			for cat in list_search:
				n_search += len(list_search[cat])
			result += str(c) + '\t' + t_max + '\t' + str(int(n_found/n_search*100)) + '\t' + '%' + '\t' + str(n_words_len) + '\t' + 'words' + '\n'
		c += 1
	g = codecs.open(os.path.join(pad_out_sent_print, project + '_' + outputfile+'_'+select_pos_text+'.txt'), 'w', 'utf-8-sig')
	text = 'Reading suggestions (2). If you read the texts in the suggested order you will encounter a maximum number of selected lexical items in a minimum text length. The first option is used in the next application, where sentences are selected to illustrate the vocabulary. Your selection codes included ' + select_pos_text +', and ' + type_se + ', which means "' + type_se_expl + '".' + '\n\n'
	g.write(text)
	g.write(result)
	g.close()
	print('result', '\n', result)
	print('total corpus', d_stats[project]['all_words'])

	## Hieronder een leuk script om teksten te zoeken die zoveel mogelijk woordenschat herhalen. Valt echter wat tegen om het zeer gespecialiseerd te maken. Waarschijnlijk beter te combineren met zinnen.

# 	av_len_text = d_stats[project]['all_words']/(len(d_stats)-1) #dit nog bekijken: zou iets teveel kunnen zijn (bedoeling is om aantal teksten in het project te vinden)

# 	list_texts_comp = []
# 	list_texts_done = []
# 	d_voc_shared = {}
# 	for c in range(1, 2):
# 		d_texts = {}
# 		d_max_shared = {}
# 		for cat in select_pos:
# 			for lem in d_lex[cat][type_se]:
# 				if lem in list_search[cat] and lem not in list_found_comp[cat]:
# 					for text in d_index[lem][cat]:
# 						if text not in list_texts_len:
# 							if text not in d_texts:
# 								d_texts[text] = [(lem, cat)]
# 							elif (lem, cat) not in d_texts[text]: #
# 								d_texts[text].append((lem, cat))
# 		for text in d_texts:
# 			list_texts_done.append(text)
# 			list_texts_done = sorted(list_texts_done)
# 		for i, text in enumerate(list_texts_done):
# 			l_len_shared = []
# 			d_len_shared = {}
# 			for text2 in list_texts_done[i+1:]:
# 				voc_shared = set(d_texts[text]).intersection(set(d_texts[text2]))
# 				d_voc_shared[(text, text2)] = voc_shared
# 				len_shared = len(voc_shared)
# 				#print(text, text2, len_shared)
# 				l_len_shared.append(len_shared)
# 				if len_shared not in d_len_shared:
# 					d_len_shared[len_shared] = [(text, text2)]
# 				else:
# 					d_len_shared[len_shared].append((text, text2))
# 			if l_len_shared:
# 				max_shared = max(l_len_shared)
# 				if max_shared > 5:
# 					if max_shared not in d_max_shared:
# 						d_max_shared[max_shared] = d_len_shared[max_shared]
# 					else:
# 						for a in d_len_shared[max_shared]:
# 							d_max_shared[max_shared].append(a)
# 		f = OrderedDict(reversed(sorted(d_max_shared.items(), key=lambda t: t[0])))
# 		for it in f:
# 			for item in f[it]:
# 				if d_stats[item[0]]['c_words'] < av_len_text*1.2 and d_stats[item[1]]['c_words'] < av_len_text*1.2:
# #					print(it, item, d_voc_shared[(item[0], item[1])]) #DIT EVENTUEEL ZICHTBAAR MAKEN
# 					pass
	return list_texts, list_texts_len, list_search

def selectsent(project, type_pr, select_level, select_spec, select_pos, list_texts, list_search): 
	print("Selecting sentences for illustrating the relevant vocabulary according to your parameters choice")
	if type_pr == 'Python_corp':
		pad_in_stat = os.path.join('..', 'output', 'stats', project)
		pad_in_sent = os.path.join('..', 'corpus', 'data', 'sent', project)
		pad_in_index = os.path.join('..', 'corpus', 'metadata', 'index_lem')
		pad_out_sent_select = os.path.join('..', 'output', 'stats', project)
		pad_out_sent_print = os.path.join('..', 'output', 'view', project)
	if type_pr == 'Local_project':
		pad_in_stat = os.path.join(project, 'metadata')
		pad_in_sent = os.path.join(project, 'data', 'sent')
		pad_out_sent_select = os.path.join(project, 'metadata')
		pad_out_sent_print = os.path.join(project, 'view')
		pad_in_index = os.path.join(project, 'metadata')
	g = codecs.open(os.path.join(pad_out_sent_select, project + '_class_lex.json'), 'r', 'utf-8-sig')
	d_lex = json.load(g)
	g.close()
	g = codecs.open(os.path.join(pad_in_index, 'corp_index_'+project + '.json'), 'r', 'utf-8-sig')
	d_index = json.load(g)
	g.close()
	g = codecs.open(os.path.join(pad_in_stat, project+'_stats.json'), 'r', 'utf-8-sig')
	d_stats = json.load(g)
	g.close()
	if select_level == 'advanced':
		out_1 = '_lev_adv'
		type_se_1 = 'adv'
		type_se_expl_1 = 'advanced vocabulary, '
	elif select_level == 'all':
		type_se_1 = 'all'
		type_se_expl_1 = 'no restrictions on learning level, '
		out_1 = 'text_all'
	if select_spec == 'high':
		out_2 = '_spec_high'
		type_se_2 = '_high'
		type_se_expl_2 = 'high specificity and high frequency'
	elif select_spec == 'medium':
		type_se_2 = '_spec_medium'
		type_se_expl_2 = 'not very frequent, but more frequent than in the reference point'
		out_2 = '_spec_medium'
	elif select_spec == 'low':
		type_se_2 = '_spec_low'
		type_se_expl_2 = 'non-typical lexemes excluded'
		out_2 = '_spec_low'
	elif select_spec == '':
		type_se_2 = '_spec_none'
		type_se_expl_2 = 'no restrictions on specificity'
		out_2 = '_spec_none'
	type_se = type_se_1+type_se_2
	type_se_expl = type_se_expl_1+type_se_expl_2
	outputfile = out_1 + out_2

	select_pos_text = '_'.join(select_pos)
	list_found = {}
	list_sent = []
	list_found_len = {}
	list_found_comp = {}
	n_words = 0
	n_words_len = 0
	d_text_lem_all = {}


	for cat in select_pos:
		list_found[cat] = []
	d_text_sent_select = {}
	for text in list_texts: #HIER ZIJN ER TWEE MOGELIJKHEDEN : ofwel list_texts ofwel list_texts_len
		lem_text = []
		lem_text_sent = []
		sent_text = {}
		list_sent_select = []
		for cat in select_pos:
			for lem in list_search[cat]:
				if lem not in list_found[cat]:
					if text in d_index[lem][cat]:
						lem_text.append((lem, cat))
						list_found[cat].append(lem)
						for sent in d_index[lem][cat][text]:
							if sent not in sent_text:
								sent_text[sent] = [lem]
							else:
								sent_text[sent].append(lem)

		f = OrderedDict(reversed(sorted(sent_text.items(), key=lambda t: len(t[1]))))
		for sent in f:
			c = 0
			for lem in f[sent]:
				if lem not in list_sent_select:
					c += 1
					list_sent_select.append(lem)

			if c != 0:
				if text not in d_text_sent_select:
					d_text_sent_select[text] = {int(sent) : []}
				else:
					d_text_sent_select[text][int(sent)] = []
				g = codecs.open(os.path.join(pad_in_sent, text + '_sent.json'), 'r', 'utf-8-sig')
				ss = json.load(g)
				n_words += ss[str(sent)]['len']
				for i, lem in enumerate(ss[str(sent)]['lem']):
					for cat in list_search:
						if lem in list_search[cat]:
							d_text_sent_select[text][int(sent)].append(i)

	print(d_text_sent_select)
	print('The estimated number of words in the selection of words is : ' + str(int(n_words*0.9)) + '. If you wish to print this in a Word file, choose export_file "yes", and run the script again.')
	return	d_text_sent_select

def selectsent_print(project, type_pr, list_texts, list_search, d_text_sent_select):
	print('Exporting the selected sentences... : see the folder "View".')
	if type_pr == 'Python_corp':
		pad_in_stat = os.path.join('..', '..', 'output', 'stats', project)
		pad_in_sent = os.path.join('..', '..', 'corpus', 'data', 'sent', project)
		pad_in_index = os.path.join('..', '..', 'corpus', 'metadata', 'index_lem')
		pad_out_sent_select = os.path.join('..', '..', 'output', 'stats', project)
		pad_out_print = os.path.join('..', '..', 'output', 'view', project)
	if type_pr == 'Local_project':
		pad_in_stat = os.path.join(project, 'metadata')
		pad_in_sent = os.path.join(project, 'data', 'sent')
		pad_out_sent_select = os.path.join(project, 'metadata')
		pad_out_print = os.path.join(project, 'view')
	document2 = Document()
	document2.add_heading('Léxico seleccionado en contexto', 1)
	o = document2.add_paragraph("")
	o.add_run("A continuación se ofrece una selección de frases que permiten consultar al menos una instancia de los elementos léxicos seleccionados. Las frases van numeradas según su posición en el texto original.").italic = True
	for text in list_texts:
		g = codecs.open(os.path.join(pad_in_sent, text+'_sent.json'), 'r', 'utf-8-sig')
		text_sent = json.load(g)
		g.close()
		if text in d_text_sent_select:
			for sent in d_text_sent_select[text]:
				p = document2.add_paragraph(str(sent) + ". ")
				for index, item in enumerate(text_sent[str(sent)]['tok']):
					if index in d_text_sent_select[text][sent]:
						p.add_run(item).underline = True
						p.add_run(" ")
					else:
						p.add_run(item +" ")
	document2.save(os.path.join(pad_out_print, project+"_sent.docx"))
